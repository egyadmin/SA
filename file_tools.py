"""
File Tools Module for Manus Clone

This module handles file operations for various file formats including:
- Text files (TXT)
- PDF files
- Word documents (DOCX)
- Excel spreadsheets (XLSX)
- AutoCAD drawings (DWG)
- Images with text extraction
"""

import os
import logging
import tempfile
from typing import Dict, List, Any, Optional, Union, BinaryIO

# Import specialized libraries for file processing
try:
    import PyPDF2
    from docx import Document
    import pandas as pd
    import ezdxf
    from PIL import Image
    import pytesseract
    import numpy as np
    import matplotlib.pyplot as plt
    import arabic_reshaper
    from bidi.algorithm import get_display
except ImportError as e:
    logging.error(f"Error importing file processing libraries: {str(e)}")
    logging.warning("Some file processing capabilities may be limited")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class FileTools:
    """File Tools class for handling various file operations"""
    
    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the File Tools with configuration
        
        Args:
            config: Dictionary containing configuration parameters
        """
        self.config = config
        self.temp_dir = config.get('TEMP_FOLDER', tempfile.gettempdir())
        
        # Ensure temp directory exists
        os.makedirs(self.temp_dir, exist_ok=True)
        
        logger.info("File Tools initialized successfully")
    
    def get_available_tools(self) -> List[str]:
        """
        Get a list of all available file tools
        
        Returns:
            List of available tool names
        """
        return [
            "read_text_file",
            "read_pdf_file",
            "read_word_file",
            "read_excel_file",
            "read_dwg_file",
            "extract_text_from_image",
            "create_text_file",
            "create_excel_file",
            "create_chart",
            "merge_pdf_files"
        ]
    
    def execute_tool(self, instruction: str) -> Dict[str, Any]:
        """
        Execute a file tool based on the instruction
        
        Args:
            instruction: Instruction describing what to do
            
        Returns:
            Dictionary containing the result of the tool execution
        """
        # Parse the instruction to determine which tool to use
        if "قراءة ملف نصي" in instruction or "read text file" in instruction.lower():
            # Extract file path from instruction
            file_path = self._extract_file_path(instruction)
            if file_path:
                return self.read_text_file(file_path)
            else:
                return {"error": "لم يتم تحديد مسار الملف"}
        
        elif "قراءة ملف pdf" in instruction or "read pdf" in instruction.lower():
            file_path = self._extract_file_path(instruction)
            if file_path:
                return self.read_pdf_file(file_path)
            else:
                return {"error": "لم يتم تحديد مسار الملف"}
        
        elif "قراءة ملف word" in instruction or "read word" in instruction.lower():
            file_path = self._extract_file_path(instruction)
            if file_path:
                return self.read_word_file(file_path)
            else:
                return {"error": "لم يتم تحديد مسار الملف"}
        
        elif "قراءة ملف excel" in instruction or "read excel" in instruction.lower():
            file_path = self._extract_file_path(instruction)
            if file_path:
                return self.read_excel_file(file_path)
            else:
                return {"error": "لم يتم تحديد مسار الملف"}
        
        elif "قراءة ملف dwg" in instruction or "read dwg" in instruction.lower():
            file_path = self._extract_file_path(instruction)
            if file_path:
                return self.read_dwg_file(file_path)
            else:
                return {"error": "لم يتم تحديد مسار الملف"}
        
        elif "استخراج نص من صورة" in instruction or "extract text from image" in instruction.lower():
            file_path = self._extract_file_path(instruction)
            if file_path:
                return self.extract_text_from_image(file_path)
            else:
                return {"error": "لم يتم تحديد مسار الملف"}
        
        elif "إنشاء ملف نصي" in instruction or "create text file" in instruction.lower():
            content = self._extract_content(instruction)
            file_path = self._extract_output_path(instruction)
            if file_path and content:
                return self.create_text_file(content, file_path)
            else:
                return {"error": "لم يتم تحديد المحتوى أو مسار الملف"}
        
        elif "إنشاء ملف excel" in instruction or "create excel file" in instruction.lower():
            data = self._extract_data(instruction)
            file_path = self._extract_output_path(instruction)
            if file_path and data:
                return self.create_excel_file(data, file_path)
            else:
                return {"error": "لم يتم تحديد البيانات أو مسار الملف"}
        
        elif "إنشاء رسم بياني" in instruction or "create chart" in instruction.lower():
            data = self._extract_data(instruction)
            chart_type = self._extract_chart_type(instruction)
            file_path = self._extract_output_path(instruction)
            if file_path and data:
                return self.create_chart(data, chart_type, file_path)
            else:
                return {"error": "لم يتم تحديد البيانات أو مسار الملف"}
        
        elif "دمج ملفات pdf" in instruction or "merge pdf" in instruction.lower():
            file_paths = self._extract_multiple_file_paths(instruction)
            output_path = self._extract_output_path(instruction)
            if file_paths and output_path:
                return self.merge_pdf_files(file_paths, output_path)
            else:
                return {"error": "لم يتم تحديد مسارات الملفات أو مسار الملف الناتج"}
        
        else:
            return {"error": "لم يتم التعرف على الأداة المطلوبة"}
    
    def read_text_file(self, file_path: str) -> Dict[str, Any]:
        """
        Read a text file and return its content
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Dictionary containing the file content
        """
        try:
            if not os.path.exists(file_path):
                return {"error": f"الملف غير موجود: {file_path}"}
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                "content": content,
                "file_path": file_path,
                "file_type": "text",
                "file_size": os.path.getsize(file_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            return {"error": f"حدث خطأ أثناء قراءة الملف النصي: {str(e)}"}
    
    def read_pdf_file(self, file_path: str) -> Dict[str, Any]:
        """
        Read a PDF file and extract its text content
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing the extracted text
        """
        try:
            if not os.path.exists(file_path):
                return {"error": f"الملف غير موجود: {file_path}"}
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                content = ""
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n\n"
            
            return {
                "content": content,
                "file_path": file_path,
                "file_type": "pdf",
                "num_pages": num_pages,
                "file_size": os.path.getsize(file_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            return {"error": f"حدث خطأ أثناء قراءة ملف PDF: {str(e)}"}
    
    def read_word_file(self, file_path: str) -> Dict[str, Any]:
        """
        Read a Word document and extract its text content
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary containing the extracted text
        """
        try:
            if not os.path.exists(file_path):
                return {"error": f"الملف غير موجود: {file_path}"}
            
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract tables if any
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "content": content,
                "tables": tables,
                "file_path": file_path,
                "file_type": "word",
                "file_size": os.path.getsize(file_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error reading Word file: {str(e)}")
            return {"error": f"حدث خطأ أثناء قراءة ملف Word: {str(e)}"}
    
    def read_excel_file(self, file_path: str) -> Dict[str, Any]:
        """
        Read an Excel spreadsheet and extract its data
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Dictionary containing the extracted data
        """
        try:
            if not os.path.exists(file_path):
                return {"error": f"الملف غير موجود: {file_path}"}
            
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            # Convert each sheet to a list of dictionaries
            sheets_data = {}
            for sheet_name, df in excel_data.items():
                sheets_data[sheet_name] = df.to_dict(orient='records')
            
            return {
                "sheets": sheets_data,
                "file_path": file_path,
                "file_type": "excel",
                "num_sheets": len(sheets_data),
                "file_size": os.path.getsize(file_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error reading Excel file: {str(e)}")
            return {"error": f"حدث خطأ أثناء قراءة ملف Excel: {str(e)}"}
    
    def read_dwg_file(self, file_path: str) -> Dict[str, Any]:
        """
        Read an AutoCAD DWG file and extract basic information
        
        Args:
            file_path: Path to the DWG file
            
        Returns:
            Dictionary containing the extracted information
        """
        try:
            if not os.path.exists(file_path):
                return {"error": f"الملف غير موجود: {file_path}"}
            
            # Open the DWG file
            doc = ezdxf.readfile(file_path)
            
            # Get model space
            msp = doc.modelspace()
            
            # Count entities by type
            entity_counts = {}
            for entity in msp:
                entity_type = entity.dxftype()
                if entity_type in entity_counts:
                    entity_counts[entity_type] += 1
                else:
                    entity_counts[entity_type] = 1
            
            # Get layers
            layers = [layer.dxf.name for layer in doc.layers]
            
            # Get drawing extents if available
            extents = None
            if hasattr(doc, 'header') and 'EXTMIN' in doc.header and 'EXTMAX' in doc.header:
                extents = {
                    'min': doc.header['EXTMIN'],
                    'max': doc.header['EXTMAX']
                }
            
            return {
                "entity_counts": entity_counts,
                "layers": layers,
                "extents": extents,
                "file_path": file_path,
                "file_type": "dwg",
                "file_size": os.path.getsize(file_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error reading DWG file: {str(e)}")
            return {"error": f"حدث خطأ أثناء قراءة ملف DWG: {str(e)}"}
    
    def extract_text_from_image(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from an image using OCR
        
        Args:
            file_path: Path to the image file
            
        Returns:
            Dictionary containing the extracted text
        """
        try:
            if not os.path.exists(file_path):
                return {"error": f"الملف غير موجود: {file_path}"}
            
            # Open the image
            image = Image.open(file_path)
            
            # Extract text using pytesseract
            text = pytesseract.image_to_string(image, lang='ara+eng')
            
            return {
                "text": text,
                "file_path": file_path,
                "file_type": "image",
                "file_size": os.path.getsize(file_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            return {"error": f"حدث خطأ أثناء استخراج النص من الصورة: {str(e)}"}
    
    def create_text_file(self, content: str, file_path: str) -> Dict[str, Any]:
        """
        Create a text file with the given content
        
        Args:
            content: Text content to write
            file_path: Path where the file should be saved
            
        Returns:
            Dictionary containing the result of the operation
        """
        try:
            # Ensure directory exists
            os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return {
                "file_path": file_path,
                "file_type": "text",
                "file_size": os.path.getsize(file_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error creating text file: {str(e)}")
            return {"error": f"حدث خطأ أثناء إنشاء الملف النصي: {str(e)}"}
    
    def create_excel_file(self, data: List[Dict[str, Any]], file_path: str) -> Dict[str, Any]:
        """
        Create an Excel file with the given data
        
        Args:
            data: List of dictionaries containing the data
            file_path: Path where the file should be saved
            
        Returns:
            Dictionary containing the result of the operation
        """
        try:
            # Ensure directory exists
            os.makedirs(os.path.dirname(os.path.abspath(file_path)), exist_ok=True)
            
            # Convert data to DataFrame
            df = pd.DataFrame(data)
            
            # Save to Excel
            df.to_excel(file_path, index=False)
            
            return {
                "file_path": file_path,
                "file_type": "excel",
                "file_size": os.path.getsize(file_path),
                "success": True
            }
        except Exception as e:
            logger.error(f"Error creating Excel file: {str(e)}")
            r
(Content truncated due to size limit. Use line ranges to read in chunks)